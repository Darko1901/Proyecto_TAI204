from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func, Date
from datetime import datetime
from io import BytesIO
from app.data.database import get_db
from app.models.pedido import Pedido
from app.models.usuario import Usuario
from app.models.detalle_pedido import DetallePedido
from app.models.producto import Producto
from app.models.inventario import Inventario
from app.security.auth import verificar_token

router = APIRouter(prefix="/v1/reportes", tags=["Reportes"])


# --- Funciones auxiliares para obtener datos ---

def _datos_ventas_por_producto(db: Session):
    return db.query(
        Producto.nombre_producto,
        func.sum(DetallePedido.cantidad).label("total_vendido"),
        func.sum(DetallePedido.cantidad * DetallePedido.precio_unitario).label("ingreso")
    ).join(DetallePedido).group_by(Producto.id_producto).all()

def _datos_pedidos_por_estado(db: Session):
    return db.query(
        Pedido.id_estado,
        func.count(Pedido.id_pedido).label("total")
    ).group_by(Pedido.id_estado).all()

def _datos_clientes(db: Session):
    return db.query(Usuario).filter(Usuario.id_rol == 1).all()

def _datos_top_productos(db: Session):
    return db.query(
        Producto.nombre_producto,
        func.sum(DetallePedido.cantidad).label("total")
    ).join(DetallePedido).group_by(Producto.id_producto).order_by(
        func.sum(DetallePedido.cantidad).desc()
    ).limit(10).all()

def _datos_ventas_tiempo(db: Session):
    # Agrupa por fecha (sin hora) y suma el total de ventas e ingresos
    return db.query(
        func.cast(Pedido.fecha_pedido, Date).label("fecha"),
        func.sum(Pedido.total).label("ingreso"),
        func.count(Pedido.id_pedido).label("cantidad")
    ).filter(Pedido.id_estado != 5).group_by(func.cast(Pedido.fecha_pedido, Date)).order_by("fecha").all()


# --- PDF ---

@router.get("/ventas/pdf")
async def reporte_ventas_pdf(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, "Reporte de Ventas por Producto - MACUIN")
    c.setFont("Helvetica", 11)
    y = 720
    for row in _datos_ventas_por_producto(db):
        c.drawString(50, y, f"{row.nombre_producto}: {row.total_vendido} uds — ${float(row.ingreso):.2f}")
        y -= 20
        if y < 50:
            c.showPage()
            y = 750
    c.save()
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": "attachment; filename=reporte_ventas.pdf"})

@router.get("/pedidos/pdf")
async def reporte_pedidos_pdf(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    estados = {1: "Recibido", 2: "Surtido", 3: "Enviado", 4: "Entregado", 5: "Cancelado"}
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, 750, "Reporte de Pedidos por Estado - MACUIN")
    c.setFont("Helvetica", 11)
    y = 720
    for row in _datos_pedidos_por_estado(db):
        nombre_estado = estados.get(row.id_estado, str(row.id_estado))
        c.drawString(50, y, f"{nombre_estado}: {row.total} pedidos")
        y -= 20
    c.save()
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": "attachment; filename=reporte_pedidos.pdf"})


# --- XLSX ---

@router.get("/ventas/xlsx")
async def reporte_ventas_xlsx(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Ventas"
    ws.append(["Producto", "Unidades Vendidas", "Ingreso Total ($)"])
    for row in _datos_ventas_por_producto(db):
        ws.append([row.nombre_producto, int(row.total_vendido), float(row.ingreso)])
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer,
                             media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=reporte_ventas.xlsx"})

@router.get("/clientes/xlsx")
async def reporte_clientes_xlsx(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Clientes"
    ws.append(["ID", "Nombre Completo", "Correo", "Teléfono", "Activo"])
    for cliente in _datos_clientes(db):
        ws.append([
            cliente.id_usuario,
            f"{cliente.nombre} {cliente.apellido_paterno} {cliente.apellido_materno}",
            cliente.correo,
            cliente.telefono or "",
            "Sí" if cliente.activo else "No"
        ])
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer,
                             media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                             headers={"Content-Disposition": "attachment; filename=reporte_clientes.xlsx"})


# --- DOCX ---

@router.get("/productos/docx")
async def reporte_top_productos_docx(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from docx import Document

    doc = Document()
    doc.add_heading("Top 10 Productos Más Vendidos - MACUIN", 0)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Producto"
    encabezados[1].text = "Unidades Vendidas"
    for row in _datos_top_productos(db):
        fila = tabla.add_row().cells
        fila[0].text = row.nombre_producto
        fila[1].text = str(int(row.total))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=reporte_productos.docx"})

@router.get("/pedidos/docx")
async def reporte_pedidos_docx(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from docx import Document

    estados = {1: "Recibido", 2: "Surtido", 3: "Enviado", 4: "Entregado", 5: "Cancelado"}
    doc = Document()
    doc.add_heading("Reporte de Pedidos por Estado - MACUIN", 0)
    for row in _datos_pedidos_por_estado(db):
        nombre_estado = estados.get(row.id_estado, str(row.id_estado))
        doc.add_paragraph(f"{nombre_estado}: {row.total} pedidos")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer,
                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=reporte_pedidos.docx"})

# --- Dashboard KPIs ---

@router.get("/dashboard")
async def obtener_kpis_dashboard(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    from datetime import date
    today = date.today()
    start_of_month = date(today.year, today.month, 1)

    # Ventas
    ventas_hoy = db.query(func.sum(Pedido.total)).filter(func.cast(Pedido.fecha_pedido, Date) == today).scalar() or 0.0
    ventas_mes = db.query(func.sum(Pedido.total)).filter(func.cast(Pedido.fecha_pedido, Date) >= start_of_month).scalar() or 0.0
    
    # Estados de Pedidos
    completados = db.query(func.count(Pedido.id_pedido)).filter(Pedido.id_estado == 4).scalar() or 0
    pendientes = db.query(func.count(Pedido.id_pedido)).filter(Pedido.id_estado < 3).scalar() or 0
    transit = db.query(func.count(Pedido.id_pedido)).filter(Pedido.id_estado == 3).scalar() or 0
    errors = db.query(func.count(Pedido.id_pedido)).filter(Pedido.id_estado == 5).scalar() or 0 # Cancelados/Errores

    # Inventario
    total_articulos = db.query(func.count(Producto.id_producto)).scalar() or 0
    
    # Stock Bajo (donde sum(inventarios.cantidad) <= inventarios.stock_minimo)
    # Para simplicidad, calculamos stock por producto primero
    stock_bajo_count = 0
    valor_total = 0.0
    productos = db.query(Producto).all()
    for p in productos:
        p_stock = sum(inv.cantidad for inv in p.inventarios)
        p_min = max((inv.stock_minimo for inv in p.inventarios), default=0)
        if p_stock <= p_min:
            stock_bajo_count += 1
        valor_total += float(p_stock * p.precio)

    return {
        "ventas_dia": float(ventas_hoy),
        "ventas_mes": float(ventas_mes),
        "pedidos_completados": completados,
        "pendientes_envio": pendientes,
        "en_transito": transit,
        "retrasos_errores": errors,
        "total_articulos": total_articulos,
        "stock_bajo": stock_bajo_count,
        "valor_inventario": valor_total,
        "hoy_formateada": datetime.now().strftime("%Y-%m-%d")
    }

@router.get("/top-productos")
async def obtener_top_productos(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    datos = _datos_top_productos(db)
    return [{"producto": row.nombre_producto, "cantidad": int(row.total)} for row in datos]

@router.get("/ventas/tendencia")
async def obtener_tendencia_ventas(db: Session = Depends(get_db), usuario_actual: Usuario = Depends(verificar_token)):
    datos = _datos_ventas_tiempo(db)
    return [{"fecha": str(row.fecha), "ingreso": float(row.ingreso), "cantidad": int(row.cantidad)} for row in datos]
